"""
Generate DOCX: Zmluva o prenájme - ROYAL STROJE s.r.o. (FO)
Faithful replica of the original PDF. Only change: product table has thick visible borders.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

ORANGE = "E87400"
BORDER = "1a1a1a"
THIN = "aaaaaa"

def shade(cell, color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'))

def set_borders(cell, sz=4, color=THIN):
    xml = (f'<w:tcBorders {nsdecls("w")}>'
           f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
           f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
           f'<w:start w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
           f'<w:end w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
           f'</w:tcBorders>')
    cell._tc.get_or_add_tcPr().append(parse_xml(xml))

def run(p, txt, sz=9, bold=False, color=None, font="Arial"):
    r = p.add_run(txt)
    r.font.size = Pt(sz); r.bold = bold; r.font.name = font
    if color: r.font.color.rgb = RGBColor(*(int(color[i:i+2],16) for i in (0,2,4)))
    return r

def fmt_cell(cell, txt, sz=8.5, bold=False, color="1a1a1a"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    run(p, txt, sz, bold, color)

def row_h(row, cm):
    row._tr.get_or_add_trPr().append(
        parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{int(cm*567)}" w:hRule="atLeast"/>'))

def orange_header(cell, txt, sz=8.5):
    shade(cell, ORANGE); set_borders(cell, 6, BORDER)
    fmt_cell(cell, txt, sz, bold=True, color="ffffff")

def label(cell, txt, sz=8):
    set_borders(cell, 4, THIN); shade(cell, "fef6ed")
    fmt_cell(cell, txt, sz, bold=True)

def val(cell, txt="", sz=8):
    set_borders(cell, 4, THIN)
    fmt_cell(cell, txt, sz)

def merge_header(row, a, b, txt, sz=8.5):
    row.cells[a].merge(row.cells[b])
    orange_header(row.cells[a], txt, sz)
    row.cells[a].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width=Cm(21); sec.page_height=Cm(29.7)
    sec.top_margin=Cm(1.2); sec.bottom_margin=Cm(0.8)
    sec.left_margin=Cm(1.5); sec.right_margin=Cm(1.5)

    # ── TITLE ──
    p = doc.add_paragraph()
    run(p, "ZMLUVA O PRENAJME HNUTELNYCH VECI \u2013 SPOTREBITELSKA ZMLUVA", 12, True, "e87400")
    run(p, " | \u00A7 663 a nasl. zak. c. 40/1964 Zb. OZ |", 7.5, False, "555555")
    p.paragraph_format.space_after = Pt(0)

    p2 = doc.add_paragraph()
    run(p2, "ROYAL STROJE s.r.o. | ICO: 57 405 425 | VPPM-FO 2026.01", 7.5, False, "888888")
    p2.paragraph_format.space_after = Pt(2); p2.paragraph_format.space_before = Pt(0)

    # orange line
    pl = doc.add_paragraph()
    pl.paragraph_format.space_before=Pt(0); pl.paragraph_format.space_after=Pt(4)
    pl._p.get_or_add_pPr().append(parse_xml(
        f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="16" w:space="1" w:color="E87400"/></w:pBdr>'))

    # ── TABLE 1: PARTIES ──
    parties = [
        ("Obchodne meno:", "ROYAL STROJE s.r.o.", "Meno a priezvisko:", ""),
        ("Sidlo:", "Recka cesta 182, 925 26 Boldog \u2013 Senec", "Adresa trvaleho bydliska:", ""),
        ("ICO / DIC / IC DPH:", "57 405 425 / 2122722063 / SK2122722063", "Datum narodenia:", ""),
        ("Zastupeny:", "Peter Krivosudsky, konatel", "Cislo OP / pasu:", ""),
        ("Tel. / E-mail:", "+421 948 555 551 / info@royalstroje.sk", "Telefon / E-mail:", ""),
        ("Zmluva c.:", "", "Fakturacny e-mail:", ""),
    ]
    t1 = doc.add_table(rows=1+len(parties), cols=4)
    merge_header(t1.rows[0], 0, 1, "PRENAJIMATEL", 9)
    merge_header(t1.rows[0], 2, 3, "NAJOMCA \u2013 SPOTREBITEL", 9)
    row_h(t1.rows[0], 0.45)
    for i,(l1,v1,l2,v2) in enumerate(parties):
        r = t1.rows[i+1]
        label(r.cells[0], l1); val(r.cells[1], v1); label(r.cells[2], l2); val(r.cells[3], v2)
        row_h(r, 0.38)
    for r in t1.rows:
        r.cells[0].width=Cm(2.8); r.cells[1].width=Cm(5.8); r.cells[2].width=Cm(3.2); r.cells[3].width=Cm(6.2)

    # ── TABLE 2: PRODUCTS (THICK BORDERS - the only change!) ──
    g = doc.add_paragraph(); g.paragraph_format.space_before=Pt(3); g.paragraph_format.space_after=Pt(0)
    t2 = doc.add_table(rows=6, cols=4)
    hdrs = ["Nazov, typ a popis predmetu prenajmu (PP)", "Vyrobne cislo", "Druh sadzby", "Sadzba/den vratane DPH (EUR)"]
    for j,h in enumerate(hdrs):
        orange_header(t2.rows[0].cells[j], h, 7.5)
        t2.rows[0].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row_h(t2.rows[0], 0.55)
    # *** THICK BLACK BORDERS on empty rows ***
    for i in range(1, 6):
        for j in range(4):
            set_borders(t2.rows[i].cells[j], sz=10, color=BORDER)
            fmt_cell(t2.rows[i].cells[j], "")
        row_h(t2.rows[i], 0.55)
    for r in t2.rows:
        r.cells[0].width=Cm(7.0); r.cells[1].width=Cm(3.5); r.cells[2].width=Cm(3.0); r.cells[3].width=Cm(4.5)

    # ── TABLE 3: INFO / FINANCE ──
    g2 = doc.add_paragraph(); g2.paragraph_format.space_before=Pt(3); g2.paragraph_format.space_after=Pt(0)
    info = [
        ("Zaciatok prenajmu:", "", "Celkove najomne vrat. DPH (EUR):", ""),
        ("Dohodnuta dlzka / koniec:", "", "Zaloha (depozit) pozadovana (EUR):", ""),
        ("Skutocny koniec prenajmu:", "", "Zaloha (depozit) vratena (EUR):", ""),
        ("Miesto pouzivania PP:", "", "Sposob platby najomneho:", ""),
        ("Miesto odovzdania PP:", "", "Sposob uhrady zalohy:", ""),
        ("Platobne podmienky:", "", "Neuctovane dni PP:", ""),
    ]
    t3 = doc.add_table(rows=1+len(info), cols=4)
    merge_header(t3.rows[0], 0, 1, "INFORMACIE O PRENAJME", 8.5)
    merge_header(t3.rows[0], 2, 3, "FINANCNE PODMIENKY", 8.5)
    row_h(t3.rows[0], 0.4)
    for i,(l1,v1,l2,v2) in enumerate(info):
        r = t3.rows[i+1]
        label(r.cells[0], l1, 7.5); val(r.cells[1], v1, 7.5); label(r.cells[2], l2, 7.5); val(r.cells[3], v2, 7.5)
        row_h(r, 0.35)
    for r in t3.rows:
        r.cells[0].width=Cm(3.5); r.cells[1].width=Cm(5.3); r.cells[2].width=Cm(4.2); r.cells[3].width=Cm(5.0)

    # ── TABLE 4: NOTES ──
    g3 = doc.add_paragraph(); g3.paragraph_format.space_before=Pt(3); g3.paragraph_format.space_after=Pt(0)
    t4 = doc.add_table(rows=2, cols=2)
    orange_header(t4.rows[0].cells[0], "Ostatne informacie o PP a prislusenstve:", 7.5)
    orange_header(t4.rows[0].cells[1], "Prevzatie PP \u2013 vyhrady k stavu (ak ziadne, nechajte prazdne):", 7.5)
    row_h(t4.rows[0], 0.35)
    val(t4.rows[1].cells[0]); val(t4.rows[1].cells[1])
    row_h(t4.rows[1], 0.8)
    for r in t4.rows:
        r.cells[0].width=Cm(9.0); r.cells[1].width=Cm(9.0)

    # ── LEGAL TEXT ──
    pl2 = doc.add_paragraph()
    pl2.paragraph_format.space_before=Pt(4); pl2.paragraph_format.space_after=Pt(4)
    run(pl2, (
        "PP je odovzdavany v stave zodpovedajucom jeho veku a opotrebeniu. Najomca (Spotrebitel) potvrdzuje "
        "prevzatie PP funkcneho a bez vyhrad (okrem uvedenych). Prilohou c. 1 su VPPM-FO 2026.01 \u2013 "
        "neoddelitelna sucast Zmluvy; v pripade rozporu ma prednost Zmluva. Najomca vyhlasuje, ze VPPM-FO "
        "si prestudoval, obdrzal ich v papierovej forme a v plnom rozsahu akceptuje. Zmluva sa spravuje zakonom "
        "c. 40/1964 Zb. OZ a zakonom c. 250/2007 Z. z. o ochrane spotrebitela. Zmluva je vyhotovena v 2 "
        "rovnopisoch. Prenajimatel spracuva osobne udaje podla cl. 6 ods. 1 pism. b) GDPR; podrobnosti na "
        "www.royalstroje.sk. Miestna prislusnost sudu sa riadi platnymi predpismi; Spotrebitel moze podat "
        "zalobu aj na sude v mieste svojho bydliska."
    ), 6.5, False, "333333")
    pl2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── TABLE 5: SIGNATURES + RETURN PROTOCOL ──
    # Original layout: left side has signatures in free-form text, right has protocol fields
    # Using 3 columns: [left signatures] | [right protocol]
    t5 = doc.add_table(rows=1, cols=2)
    set_borders(t5.rows[0].cells[0], 6, BORDER)
    set_borders(t5.rows[0].cells[1], 6, BORDER)

    # Left cell - signatures
    left = t5.rows[0].cells[0]
    shade(left, ORANGE)
    left.text = ""
    p = left.paragraphs[0]
    run(p, "PODPISY ZMLUVNYCH STRAN  |", 8.5, True, "ffffff")
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)

    # Right cell - protocol header
    right = t5.rows[0].cells[1]
    shade(right, ORANGE)
    right.text = ""
    p = right.paragraphs[0]
    run(p, "PROTOKOL O VRATENI PP", 8.5, True, "ffffff")
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)

    left.width = Cm(9.0); right.width = Cm(9.0)
    row_h(t5.rows[0], 0.4)

    # Signature body as separate table below
    t6 = doc.add_table(rows=1, cols=2)
    set_borders(t6.rows[0].cells[0], 4, THIN)
    set_borders(t6.rows[0].cells[1], 4, THIN)

    # Left: signatures
    lc = t6.rows[0].cells[0]
    lc.text = ""
    lines_left = [
        ("", False), ("Za prenajimatela:", True), ("Meno a priezvisko:", False),
        ("", False), ("_________________________", False),
        ("Datum:", False), ("", False), ("_________________________", False),
        ("Podpis:", False), ("", False),
        ("Najomca \u2013 spotrebitel:", True), ("Meno a priezvisko:", False),
        ("", False), ("_________________________", False),
        ("Datum:", False), ("", False), ("_________________________", False),
        ("Podpis:", False),
    ]
    first = True
    for txt, bold in lines_left:
        if first:
            p = lc.paragraphs[0]; first = False
        else:
            p = lc.add_paragraph()
        p.paragraph_format.space_before = Pt(0.5); p.paragraph_format.space_after = Pt(0.5)
        run(p, txt, 7.5 if not bold else 8, bold, "1a1a1a")
    lc.width = Cm(9.0)

    # Right: return protocol
    rc = t6.rows[0].cells[1]
    rc.text = ""
    lines_right = [
        ("Datum a cas vratenia:", True),
        ("_________________________________________", False),
        ("Stav PP pri vrateni:", True),
        ("_________________________________________", False),
        ("Poskodenia / chybajuce prislusenstvo:", True),
        ("_________________________________________", False),
        ("Vycisteny:  \u2610 Ano  \u2610 Nie     Fotodokumentacia:  \u2610 Ano  \u2610 Nie", False),
        ("Podpis prenajimatela:", True),
        ("_________________________________________", False),
        ("Podpis najomcu:", True),
        ("_________________________________________", False),
    ]
    first = True
    for txt, bold in lines_right:
        if first:
            p = rc.paragraphs[0]; first = False
        else:
            p = rc.add_paragraph()
        p.paragraph_format.space_before = Pt(0.5); p.paragraph_format.space_after = Pt(0.5)
        run(p, txt, 7.5 if not bold else 8, bold, "1a1a1a")
    rc.width = Cm(9.0)

    # Save
    out = os.path.join(os.path.dirname(__file__), "zmluva-fo-royal-stroje.docx")
    doc.save(out)
    print(f"OK: {out}")

if __name__ == "__main__":
    main()
